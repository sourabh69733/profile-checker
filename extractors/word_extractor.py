# extractors/word_extractor.py
import os
import subprocess
import tempfile
from docx import Document
from .base import BaseExtractor

class WordExtractor(BaseExtractor):
    """Extract text from DOCX and DOC files"""
    
    supported_extensions = ['.docx', '.doc']
    
    def __init__(self, libreoffice_path: str = None):
        self.libreoffice_path = libreoffice_path or self._find_libreoffice()
    
    def extract(self, file_path: str) -> str:
        self.validate_file(file_path)
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return self._extract_docx(file_path)
        elif ext == '.doc':
            return self._extract_doc(file_path)
        else:
            raise ValueError(f"Unsupported format: {ext}")
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(table_text)
        
        # Extract headers/footers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                if para.text.strip():
                    text_parts.insert(0, para.text)
        
        return "\n".join(text_parts)
    
    def _extract_table(self, table) -> str:
        """Extract text from Word table"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))
        return "\n".join(rows)
    
    def _extract_doc(self, file_path: str) -> str:
        """Extract text from legacy DOC files by converting to DOCX"""
        if not self.libreoffice_path:
            raise RuntimeError(
                "LibreOffice not found. Please install it to process .doc files"
            )
        
        with tempfile.TemporaryDirectory() as temp_dir:
            # Convert DOC to DOCX
            result = subprocess.run([
                self.libreoffice_path,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', temp_dir,
                file_path
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode != 0:
                raise RuntimeError(f"Failed to convert DOC file: {result.stderr}")
            
            # Find converted file
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            docx_path = os.path.join(temp_dir, f"{base_name}.docx")
            
            if not os.path.exists(docx_path):
                raise RuntimeError("DOC conversion failed - output file not found")
            
            return self._extract_docx(docx_path)
    
    def _find_libreoffice(self) -> str:
        """Auto-detect LibreOffice installation"""
        possible_paths = [
            '/usr/bin/libreoffice',
            '/usr/bin/soffice',
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
            'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        # Try to find in PATH
        try:
            result = subprocess.run(['which', 'libreoffice'], capture_output=True, text=True)
            if result.returncode == 0:
                return result.stdout.strip()
        except:
            pass
        
        return None