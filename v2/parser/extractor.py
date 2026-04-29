"""
Extract text from various file formats (PDF, DOCX, TXT)
"""
import io
import re
from pathlib import Path
from typing import Optional, Tuple
import fitz  # PyMuPDF
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract
from pdfminer.layout import LAParams


class ResumeExtractor:
    """Extract text from resume files with multiple fallback methods."""
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt', '.rtf']
    
    def __init__(self):
        self.extraction_method = None
        self.confidence = 1.0
    
    def extract(self, file_path: str | Path) -> Tuple[str, dict]:
        """
        Extract text from file.
        
        Returns:
            Tuple of (extracted_text, metadata)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            text = self._extract_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            text = self._extract_docx(file_path)
        elif suffix == '.txt':
            text = self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported format: {suffix}")
        
        # Clean extracted text
        text = self._clean_text(text)
        
        metadata = {
            "file_name": file_path.name,
            "file_type": suffix,
            "extraction_method": self.extraction_method,
            "char_count": len(text),
            "word_count": len(text.split()),
            "confidence": self.confidence
        }
        
        return text, metadata
    
    def extract_from_bytes(self, file_bytes: bytes, file_type: str) -> Tuple[str, dict]:
        """Extract text from file bytes."""
        file_type = file_type.lower()
        if not file_type.startswith('.'):
            file_type = f'.{file_type}'
        
        if file_type == '.pdf':
            text = self._extract_pdf_bytes(file_bytes)
        elif file_type in ['.docx', '.doc']:
            text = self._extract_docx_bytes(file_bytes)
        elif file_type == '.txt':
            text = file_bytes.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported format: {file_type}")
        
        text = self._clean_text(text)
        
        metadata = {
            "file_type": file_type,
            "extraction_method": self.extraction_method,
            "char_count": len(text),
            "word_count": len(text.split()),
            "confidence": self.confidence
        }
        
        return text, metadata
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF with multiple methods."""
        text = ""
        
        # Method 1: PyMuPDF (fast, good for most PDFs)
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text("text")
            doc.close()
            
            if self._is_valid_extraction(text):
                self.extraction_method = "pymupdf"
                self.confidence = 0.95
                return text
        except Exception as e:
            pass
        
        # Method 2: PDFMiner (better for complex layouts)
        try:
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5
            )
            text = pdfminer_extract(str(file_path), laparams=laparams)
            
            if self._is_valid_extraction(text):
                self.extraction_method = "pdfminer"
                self.confidence = 0.90
                return text
        except Exception as e:
            pass
        
        # Method 3: PyMuPDF with different settings
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text("blocks")
            doc.close()
            self.extraction_method = "pymupdf_blocks"
            self.confidence = 0.80
            return str(text)
        except Exception as e:
            pass
        
        self.extraction_method = "failed"
        self.confidence = 0.0
        return ""
    
    def _extract_pdf_bytes(self, file_bytes: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text("text")
            doc.close()
            self.extraction_method = "pymupdf"
            self.confidence = 0.95
            return text
        except Exception:
            # Fallback to pdfminer
            try:
                text = pdfminer_extract(io.BytesIO(file_bytes))
                self.extraction_method = "pdfminer"
                self.confidence = 0.90
                return text
            except Exception:
                self.extraction_method = "failed"
                self.confidence = 0.0
                return ""
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            self.extraction_method = "python-docx"
            self.confidence = 0.95
            return "\n".join(paragraphs)
        except Exception as e:
            self.extraction_method = "failed"
            self.confidence = 0.0
            return ""
    
    def _extract_docx_bytes(self, file_bytes: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            self.extraction_method = "python-docx"
            self.confidence = 0.95
            return "\n".join(paragraphs)
        except Exception:
            self.extraction_method = "failed"
            self.confidence = 0.0
            return ""
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                self.extraction_method = f"txt_{encoding}"
                self.confidence = 1.0
                return text
            except UnicodeDecodeError:
                continue
        
        self.extraction_method = "failed"
        self.confidence = 0.0
        return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        if not text:
            return ""
        
        # Remove null bytes and control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize whitespace
        text = re.sub(r'\t', ' ', text)
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove page numbers
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        
        return text.strip()
    
    def _is_valid_extraction(self, text: str) -> bool:
        """Check if extraction produced valid text."""
        if not text or len(text) < 100:
            return False
        
        # Check for minimum word count
        words = text.split()
        if len(words) < 20:
            return False
        
        # Check for reasonable character ratio (not all special chars)
        alpha_count = sum(1 for c in text if c.isalpha())
        if alpha_count / len(text) < 0.5:
            return False
        
        return True